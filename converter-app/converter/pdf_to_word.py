"""
PDF to Word converter engine.

Primary: pdf2docx for layout-preserving conversion
Fallback: PyMuPDF extraction + python-docx reconstruction
OCR: Tesseract for scanned PDFs
"""

import os
import tempfile
import fitz  # PyMuPDF
from pdf2docx import Converter as Pdf2DocxConverter

from .ocr_engine import is_scanned_pdf, ocr_pdf_to_text, ocr_pdf_page_to_image
from .layout_detect import detect_layout, has_tables
from .utils import get_output_path


def convert_pdf_to_word(
    pdf_path: str,
    output_path: str = "",
    progress_callback=None,
    ocr_fallback: bool = True,
) -> str:
    """
    Convert PDF to Word (.docx) preserving structure.

    Strategy:
    1. Check if PDF is scanned → use OCR path
    2. Digital PDF → use pdf2docx for layout reconstruction
    3. If pdf2docx fails → fall back to PyMuPDF extraction

    Args:
        pdf_path: Path to input PDF file
        output_path: Optional output path (auto-generated if empty)
        progress_callback: Optional callback(percent: int)
        ocr_fallback: Whether to use OCR for scanned PDFs

    Returns:
        Path to the generated .docx file
    """
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    if not output_path:
        output_path = get_output_path(pdf_path, ".docx")

    # Validate PDF
    try:
        doc = fitz.open(pdf_path)
        page_count = len(doc)
        doc.close()
    except Exception as e:
        raise ValueError(f"Cannot open PDF (may be corrupted): {e}")

    if progress_callback:
        progress_callback(5)

    # Step 1: Detect if scanned
    scanned = is_scanned_pdf(pdf_path)

    if progress_callback:
        progress_callback(10)

    if scanned and ocr_fallback:
        return _convert_scanned_pdf(pdf_path, output_path, progress_callback)
    else:
        return _convert_digital_pdf(pdf_path, output_path, progress_callback)


def _convert_digital_pdf(pdf_path: str, output_path: str, progress_callback=None) -> str:
    """
    Convert a digital (text-based) PDF to Word using pdf2docx.
    This preserves layout, tables, borders, images, fonts, spacing, and page breaks.
    """
    try:
        if progress_callback:
            progress_callback(15)

        cv = Pdf2DocxConverter(pdf_path)

        if progress_callback:
            progress_callback(25)

        # Convert all pages
        cv.convert(output_path)

        if progress_callback:
            progress_callback(90)

        cv.close()

        if progress_callback:
            progress_callback(100)

        return output_path

    except Exception as e:
        # Fallback to PyMuPDF extraction
        if progress_callback:
            progress_callback(30)
        return _fallback_pymupdf_convert(pdf_path, output_path, progress_callback)


def _convert_scanned_pdf(pdf_path: str, output_path: str, progress_callback=None) -> str:
    """
    Convert a scanned PDF to Word using OCR + layout detection.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import io

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)

    for page_idx in range(total_pages):
        if progress_callback:
            pct = 15 + int((page_idx / total_pages) * 75)
            progress_callback(pct)

        # Get page as image for layout detection
        page_img = ocr_pdf_page_to_image(pdf_path, page_idx, dpi=300)

        # Detect layout regions
        regions = detect_layout(page_img)

        # OCR the page
        page = pdf_doc[page_idx]
        mat = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_data = pix.tobytes("png")

        from PIL import Image
        import pytesseract
        from .utils import get_tesseract_path

        tess_path = get_tesseract_path()
        if tess_path:
            pytesseract.pytesseract.tesseract_cmd = tess_path

        img = Image.open(io.BytesIO(img_data))

        # Process each region
        for region in regions:
            rtype = region["type"]
            bbox = region["bbox"]

            # Crop region from image
            cropped = img.crop(bbox)

            if rtype == "figure":
                # Insert as image
                img_stream = io.BytesIO()
                cropped.save(img_stream, format="PNG")
                img_stream.seek(0)

                # Calculate size relative to page
                page_width = img.width
                region_width = bbox[2] - bbox[0]
                doc_width_inches = 6.0  # usable width
                img_width = doc_width_inches * (region_width / page_width)
                img_width = min(img_width, 6.0)

                doc.add_picture(img_stream, width=Inches(img_width))

            elif rtype == "table":
                # OCR the table region and try to structure it
                table_text = pytesseract.image_to_string(cropped, lang="eng+ron")
                lines = [l.strip() for l in table_text.split("\n") if l.strip()]

                if lines:
                    # Try to detect columns by consistent spacing
                    rows_data = []
                    for line in lines:
                        cells = [c.strip() for c in line.split("  ") if c.strip()]
                        if not cells:
                            cells = [line]
                        rows_data.append(cells)

                    # Normalize column count
                    max_cols = max(len(r) for r in rows_data) if rows_data else 1
                    for row in rows_data:
                        while len(row) < max_cols:
                            row.append("")

                    if max_cols > 0 and len(rows_data) > 0:
                        table = doc.add_table(rows=len(rows_data), cols=max_cols)
                        table.style = "Table Grid"
                        for ri, row in enumerate(rows_data):
                            for ci, cell in enumerate(row):
                                table.rows[ri].cells[ci].text = cell
                    else:
                        doc.add_paragraph(table_text)
                else:
                    # Empty table region — insert as image
                    img_stream = io.BytesIO()
                    cropped.save(img_stream, format="PNG")
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(5.0))

            elif rtype == "title":
                text = pytesseract.image_to_string(cropped, lang="eng+ron").strip()
                if text:
                    doc.add_heading(text, level=1)

            else:
                # Regular text
                text = pytesseract.image_to_string(cropped, lang="eng+ron").strip()
                if text:
                    doc.add_paragraph(text)

        # Page break between pages (except last)
        if page_idx < total_pages - 1:
            doc.add_page_break()

    pdf_doc.close()

    doc.save(output_path)

    if progress_callback:
        progress_callback(100)

    return output_path


def _fallback_pymupdf_convert(pdf_path: str, output_path: str, progress_callback=None) -> str:
    """
    Fallback converter using PyMuPDF for extraction + python-docx for output.
    Used when pdf2docx fails.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)

    for page_idx in range(total_pages):
        if progress_callback:
            pct = 35 + int((page_idx / total_pages) * 55)
            progress_callback(pct)

        page = pdf_doc[page_idx]

        # Extract text blocks with position info
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

        for block in blocks:
            if block["type"] == 0:  # Text block
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue

                    para = doc.add_paragraph()

                    for span in spans:
                        text = span.get("text", "")
                        if not text:
                            continue

                        run = para.add_run(text)
                        size = span.get("size", 11)
                        run.font.size = Pt(max(6, min(72, size)))

                        # Font name
                        font_name = span.get("font", "")
                        if font_name:
                            # Clean up font name
                            clean_name = font_name.split("+")[-1] if "+" in font_name else font_name
                            clean_name = clean_name.split("-")[0] if "-" in clean_name else clean_name
                            run.font.name = clean_name

                        # Bold detection
                        flags = span.get("flags", 0)
                        if flags & 2**4:  # Bold
                            run.font.bold = True
                        if flags & 2**1:  # Italic
                            run.font.italic = True

                        # Color
                        color = span.get("color", 0)
                        if color and color != 0:
                            r = (color >> 16) & 0xFF
                            g = (color >> 8) & 0xFF
                            b = color & 0xFF
                            if (r, g, b) != (0, 0, 0):
                                run.font.color.rgb = RGBColor(r, g, b)

            elif block["type"] == 1:  # Image block
                try:
                    img_data = block.get("image", b"")
                    if img_data:
                        img_stream = io.BytesIO(img_data)
                        # Calculate width
                        bbox = block.get("bbox", (0, 0, 100, 100))
                        img_width_pts = bbox[2] - bbox[0]
                        page_width_pts = page.rect.width
                        img_width_inches = min(6.0, (img_width_pts / page_width_pts) * 6.5)
                        doc.add_picture(img_stream, width=Inches(img_width_inches))
                except Exception:
                    pass

        # Page break between pages
        if page_idx < total_pages - 1:
            doc.add_page_break()

    pdf_doc.close()
    doc.save(output_path)

    if progress_callback:
        progress_callback(100)

    return output_path
