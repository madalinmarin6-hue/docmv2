"""
Word to PDF converter engine.

Primary: LibreOffice headless for perfect rendering
Fallback: PyMuPDF-based conversion from extracted content
"""

import os
import subprocess
import tempfile
import shutil

from .utils import get_libreoffice_path, get_output_path


def convert_word_to_pdf(
    docx_path: str,
    output_path: str = "",
    progress_callback=None,
) -> str:
    """
    Convert Word (.docx) to PDF preserving 100% structure.

    Uses LibreOffice headless mode for perfect rendering of:
    - Tables, borders, boxes
    - Images, fonts, spacing
    - Page breaks, columns
    - Headers and footers

    Args:
        docx_path: Path to input .docx file
        output_path: Optional output path (auto-generated if empty)
        progress_callback: Optional callback(percent: int)

    Returns:
        Path to the generated .pdf file
    """
    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"Word file not found: {docx_path}")

    if not output_path:
        output_path = get_output_path(docx_path, ".pdf")

    if progress_callback:
        progress_callback(10)

    # Try LibreOffice first (best quality)
    lo_path = get_libreoffice_path()

    if lo_path:
        return _convert_with_libreoffice(docx_path, output_path, lo_path, progress_callback)
    else:
        return _convert_fallback(docx_path, output_path, progress_callback)


def _convert_with_libreoffice(
    docx_path: str,
    output_path: str,
    lo_path: str,
    progress_callback=None,
) -> str:
    """
    Convert using LibreOffice headless mode.
    This produces the highest quality PDF output.
    """
    if progress_callback:
        progress_callback(20)

    # LibreOffice requires an output directory, not a file path
    output_dir = os.path.dirname(output_path) or os.getcwd()
    base_name = os.path.splitext(os.path.basename(docx_path))[0]

    # Use a temp directory to avoid conflicts
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Copy source file to temp dir to avoid path issues
        tmp_input = os.path.join(tmp_dir, os.path.basename(docx_path))
        shutil.copy2(docx_path, tmp_input)

        if progress_callback:
            progress_callback(30)

        # Run LibreOffice headless
        cmd = [
            lo_path,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", tmp_dir,
            tmp_input,
        ]

        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,  # 2 minute timeout
                cwd=tmp_dir,
            )

            if progress_callback:
                progress_callback(80)

            # Find the generated PDF
            expected_pdf = os.path.join(tmp_dir, base_name + ".pdf")
            # Also check with original filename base
            orig_base = os.path.splitext(os.path.basename(docx_path))[0]
            alt_pdf = os.path.join(tmp_dir, orig_base + ".pdf")

            pdf_found = None
            if os.path.isfile(expected_pdf):
                pdf_found = expected_pdf
            elif os.path.isfile(alt_pdf):
                pdf_found = alt_pdf
            else:
                # Search for any PDF in temp dir
                for f in os.listdir(tmp_dir):
                    if f.lower().endswith(".pdf"):
                        pdf_found = os.path.join(tmp_dir, f)
                        break

            if pdf_found:
                shutil.copy2(pdf_found, output_path)

                if progress_callback:
                    progress_callback(100)

                return output_path
            else:
                raise RuntimeError(
                    f"LibreOffice did not produce a PDF.\n"
                    f"stdout: {result.stdout}\n"
                    f"stderr: {result.stderr}"
                )

        except subprocess.TimeoutExpired:
            raise RuntimeError("LibreOffice conversion timed out (120s)")

        except FileNotFoundError:
            raise RuntimeError(f"LibreOffice not found at: {lo_path}")


def _convert_fallback(
    docx_path: str,
    output_path: str,
    progress_callback=None,
) -> str:
    """
    Fallback conversion using python-docx for reading + reportlab/fitz for PDF.
    Quality is lower than LibreOffice but works without external dependencies.
    """
    import fitz  # PyMuPDF

    if progress_callback:
        progress_callback(20)

    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        raise RuntimeError("python-docx is required for fallback conversion")

    doc = Document(docx_path)

    if progress_callback:
        progress_callback(30)

    # Create PDF with PyMuPDF
    pdf_doc = fitz.open()

    # Page dimensions (A4)
    page_width = 595.28
    page_height = 841.89
    margin = 72  # 1 inch

    current_y = margin
    page = pdf_doc.new_page(width=page_width, height=page_height)

    usable_width = page_width - 2 * margin
    total_paragraphs = len(doc.paragraphs)

    for para_idx, paragraph in enumerate(doc.paragraphs):
        if progress_callback and total_paragraphs > 0:
            pct = 30 + int((para_idx / total_paragraphs) * 60)
            progress_callback(pct)

        text = paragraph.text
        if not text.strip():
            current_y += 12
            continue

        # Determine font size from style
        font_size = 11
        is_bold = False
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")

        if is_heading:
            level = 1
            try:
                level = int(paragraph.style.name.split()[-1])
            except (ValueError, IndexError):
                pass
            font_size = max(24 - (level - 1) * 4, 12)
            is_bold = True

        # Check runs for formatting
        for run in paragraph.runs:
            if run.font.size:
                font_size = run.font.size.pt
            if run.font.bold:
                is_bold = True

        line_height = font_size * 1.4

        # Check page overflow
        if current_y + line_height > page_height - margin:
            page = pdf_doc.new_page(width=page_width, height=page_height)
            current_y = margin

        # Insert text
        fontname = "helv" if not is_bold else "hebo"

        try:
            text_writer = fitz.TextWriter(page.rect)
            text_writer.append(
                (margin, current_y),
                text,
                fontsize=font_size,
                font=fitz.Font(fontname),
            )
            text_writer.write_text(page)
        except Exception:
            # Simple fallback
            page.insert_text(
                (margin, current_y),
                text[:100],  # Truncate long lines
                fontsize=font_size,
            )

        current_y += line_height

    # Process tables
    for table in doc.tables:
        if current_y + 50 > page_height - margin:
            page = pdf_doc.new_page(width=page_width, height=page_height)
            current_y = margin

        row_height = 20
        col_width = usable_width / max(len(table.columns), 1)

        for row in table.rows:
            if current_y + row_height > page_height - margin:
                page = pdf_doc.new_page(width=page_width, height=page_height)
                current_y = margin

            for ci, cell in enumerate(row.cells):
                x = margin + ci * col_width
                # Draw cell border
                rect = fitz.Rect(x, current_y, x + col_width, current_y + row_height)
                page.draw_rect(rect, color=(0, 0, 0), width=0.5)
                # Insert cell text
                cell_text = cell.text[:30] if cell.text else ""
                try:
                    page.insert_text((x + 3, current_y + 14), cell_text, fontsize=9)
                except Exception:
                    pass

            current_y += row_height

    if progress_callback:
        progress_callback(95)

    pdf_doc.save(output_path)
    pdf_doc.close()

    if progress_callback:
        progress_callback(100)

    return output_path
